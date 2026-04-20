"""
core/exporter.py
Esportazione del testo trascritto in TXT e DOCX con formattazione.
"""

import os
from pathlib import Path


def export_txt(html_text: str, output_path: str) -> str:
    """
    Salva il testo in formato TXT (testo semplice, senza tag HTML).
    Restituisce il percorso del file salvato.
    """
    from PyQt5.QtGui import QTextDocument
    doc = QTextDocument()
    doc.setHtml(html_text)
    plain_text = doc.toPlainText()

    output_path = _ensure_extension(output_path, ".txt")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(plain_text)
    return output_path


def export_docx(html_text: str, output_path: str) -> str:
    """
    Salva il testo in formato DOCX preservando la formattazione
    (grassetto, corsivo, sottolineato, allineamento, dimensione font).
    Restituisce il percorso del file salvato.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PyQt5.QtGui import QTextDocument, QTextBlock, QTextCharFormat, QTextBlockFormat
    from PyQt5.QtCore import Qt

    qt_doc = QTextDocument()
    qt_doc.setHtml(html_text)

    word_doc = Document()
    # Margini ridotti per un look più pulito
    for section in word_doc.sections:
        section.top_margin = _cm(2)
        section.bottom_margin = _cm(2)
        section.left_margin = _cm(2.5)
        section.right_margin = _cm(2.5)

    block = qt_doc.begin()
    while block.isValid():
        block_fmt: QTextBlockFormat = block.blockFormat()
        paragraph = word_doc.add_paragraph()

        # Allineamento
        alignment = block_fmt.alignment()
        if alignment & Qt.AlignRight:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment & Qt.AlignHCenter:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment & Qt.AlignJustify:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        it = block.begin()
        while not it.atEnd():
            fragment = it.fragment()
            if fragment.isValid():
                char_fmt: QTextCharFormat = fragment.charFormat()
                text = fragment.text()
                if not text:
                    it += 1
                    continue

                run = paragraph.add_run(text)

                # Grassetto
                run.bold = char_fmt.fontWeight() >= 700
                # Corsivo
                run.italic = char_fmt.fontItalic()
                # Sottolineato
                run.underline = char_fmt.fontUnderline()

                # Dimensione font
                font_size = char_fmt.fontPointSize()
                if font_size > 0:
                    run.font.size = Pt(font_size)

                # Colore testo (ignora nero/default)
                color = char_fmt.foreground().color()
                if color.isValid() and color.name() not in ("#000000", "#000", ""):
                    run.font.color.rgb = RGBColor(color.red(), color.green(), color.blue())

            it += 1

        block = block.next()

    output_path = _ensure_extension(output_path, ".docx")
    word_doc.save(output_path)
    return output_path


def _ensure_extension(path: str, ext: str) -> str:
    p = Path(path)
    if p.suffix.lower() != ext:
        p = p.with_suffix(ext)
    return str(p)


def _cm(value: float):
    from docx.shared import Cm
    return Cm(value)
